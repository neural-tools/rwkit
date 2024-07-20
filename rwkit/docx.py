"""
Docx file I/O
"""

from pathlib import Path
from typing import List, Union

try:
    import docx

    _HAVE_DOCX = True
except ImportError:
    _HAVE_DOCX = False


def read_docx(
    filename: Union[str, Path], *, skip_empty_lines: bool = True
) -> List[str]:
    """
    Read docx file and return lines as list of strings.

    Args:
        filename (Union[str, Path]): File to read.
        skip_empty_lines (bool, optional): Whether empty lines should be skipped.
            Defaults to True.

    Returns:
        List[str]: Lines from file.

    Raises:
        ModuleNotFoundError: If package 'python-docx' is not installed.
    """
    if not _HAVE_DOCX:
        raise ModuleNotFoundError(
            "No module named 'docx'. Install with: $ pip install python-docx"
        )

    doc = docx.Document(docx=str(filename))

    # Iterate over all paragraphs
    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    txts = content.rstrip("\n").split("\n")

    if skip_empty_lines:
        txts = [txt for txt in txts if len(txt) > 0]

    return txts


def write_docx(filename: Union[str, Path], text: Union[str, List[str]]) -> None:
    """
    Write strings to a docx file.

    Args:
        filename (Union[str, Path]): File to write to.
        text (Union[str, List[str]]): String(s) to write. If a string, it is written
            as a single paragraph. If a list of strings, each string is written as a
            separate paragraph.

    Raises:
        ModuleNotFoundError: If package 'python-docx' is not installed.
        TypeError: If text is not a string or list of strings.
    """
    if not _HAVE_DOCX:
        raise ModuleNotFoundError(
            "No module named 'docx'. Install with: $ pip install python-docx"
        )

    if isinstance(text, str):
        txts = [text]
    elif isinstance(text, list):
        txts = text
    else:
        raise TypeError("Expected str or list, got %s" % type(text))

    doc = docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("\n".join(txts) + "\n")
    doc.save(path_or_stream=filename)
