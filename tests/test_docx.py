"""
Test docx file I/O
"""

import itertools
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

import docx

from rwkit.docx import read_docx, write_docx


class TestDocx(unittest.TestCase):
    """TestDocx"""

    def test_read_docx(self):
        """test_read_docx"""

        txts_expected_list = [
            ["These", "are words", "of a sentence"],
            [
                "",
                "These",
                "",
                "are words",
                "",
                "",
                "of a sentence",
            ],
        ]

        skip_empty_lines_list = [True, False]

        for txts_expected, skip_empty_lines in itertools.product(
            txts_expected_list, skip_empty_lines_list
        ):
            with TemporaryDirectory() as tmpdir:
                filepath = Path(tmpdir) / "file"

                # Write to file
                doc = docx.Document()
                paragraph = doc.add_paragraph()
                paragraph.add_run("\n".join(txts_expected) + "\n")
                doc.save(path_or_stream=filepath)

                # Read file contents
                txts_observed = read_docx(
                    filename=filepath, skip_empty_lines=skip_empty_lines
                )

                if skip_empty_lines:
                    txts_expected = [txt for txt in txts_expected if len(txt) > 0]

                self.assertEqual(
                    txts_expected,
                    txts_observed,
                    "read_docx() failed.\n"
                    "Parameters:\n"
                    f"  skip_empty_lines: {skip_empty_lines}\n"
                    f"Expected: '{txts_expected}'\n"
                    f"Observed: '{txts_observed}'",
                )

    def test_write_docx(self):
        """test_write_docx"""

        txts_expected_list = [
            ["These", "are words", "of a sentence"],
            [
                "",
                "These",
                "",
                "are words",
                "",
                "",
                "of a sentence",
            ],
        ]

        for txts_expected in txts_expected_list:
            with TemporaryDirectory() as tmpdir:
                filepath = Path(tmpdir) / "file"

                # Write to file
                write_docx(
                    filename=filepath,
                    text=txts_expected,
                )

                # Read file contents
                doc = docx.Document(docx=filepath)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                txts_observed = content.rstrip("\n").split("\n")

                self.assertEqual(
                    txts_expected,
                    txts_observed,
                    "write_docx() failed.\n"
                    f"Expected: '{txts_expected}'\n"
                    f"Observed: '{txts_observed}'",
                )

    def test_read_write_docx(self):
        """test_read_write_docx"""

        txts_expected_list = [
            ["These", "are words", "of a sentence"],
            [
                "",
                "These",
                "",
                "are words",
                "",
                "",
                "of a sentence",
            ],
        ]
        skip_empty_lines_list = [True, False]

        for txts_expected, skip_empty_lines in itertools.product(
            txts_expected_list, skip_empty_lines_list
        ):
            with TemporaryDirectory() as tmpdir:
                filepath = Path(tmpdir) / "file"

                # Write to file
                write_docx(
                    filename=filepath,
                    text=txts_expected,
                )

                # Read file contents
                txts_observed = read_docx(
                    filename=filepath, skip_empty_lines=skip_empty_lines
                )

                if skip_empty_lines:
                    txts_expected = [txt for txt in txts_expected if len(txt) > 0]

                self.assertEqual(
                    txts_expected,
                    txts_observed,
                    "read_docx() failed.\n"
                    "Parameters:\n"
                    f"  skip_empty_lines: {skip_empty_lines}\n"
                    f"Expected: '{txts_expected}'\n"
                    f"Observed: '{txts_observed}'",
                )
